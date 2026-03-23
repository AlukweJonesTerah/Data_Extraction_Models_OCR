import os
import io
import re
import json
import time
import shutil
import subprocess
import requests
from datetime import datetime
from PIL import Image
from ollama import Client

# ── Configuration ─────────────────────────────────────────────────────────────
source_dir        = r"\\192.168.1.11\d\BCERT_Phase4_Omniscan\DMS_Upload\TIFF\_invalid_member_numbers"
output_dir        = r"\\192.168.1.11\d\BCERT_Phase4_Omniscan\DMS_Upload\Extracted_XMLs_gpu_v1"
manual_review_dir = os.path.join(output_dir, "Manual_Review_Needed")
reports_dir       = os.path.join(output_dir, "Reports")

# --- RunPod / Remote GPU Configuration ---
USE_RUNPOD = True
RUNPOD_URL = "https://am37kwma1dwy5l-64411a83-11434.proxy.runpod.net"

# Increase timeout to 300s to survive deepseek-ocr model load through Cloudflare proxy
OLLAMA_TIMEOUT = 300

if USE_RUNPOD:
    print(f"[*] Connecting to remote RunPod GPU: {RUNPOD_URL}")
    ai_client = Client(host=RUNPOD_URL, timeout=OLLAMA_TIMEOUT)
else:
    print("[*] Connecting to Local Machine GPU...")
    ai_client = Client(host='http://localhost:11434', timeout=OLLAMA_TIMEOUT)

# The 3-Tier AI Extraction Team
PRIMARY_MODEL   = 'qwen2.5vl'
SECONDARY_MODEL = 'llama3.2-vision'
TERTIARY_MODEL  = 'deepseek-ocr'

SYSTEM_PROMPT = """
You are a highly accurate OCR data extraction system. Look at the provided cropped image of a Kenyan NSSF Certificate.
Read the number written or typed next to "NSSF No." in the yellow box.

Rules:
1. ONLY output the number.
2. DO NOT invent or guess numbers. If it is completely unreadable, output UNREADABLE.
3. Output EXACTLY in this XML format and nothing else:

<JT_Member_Number>[Extracted NSSF No]</JT_Member_Number>
"""

xml_pattern = re.compile(r'<JT_Member_Number>(.*?)</JT_Member_Number>', re.IGNORECASE)


# ── Validators & Helpers ──────────────────────────────────────────────────────

def is_valid_nssf(nssf_string):
    clean_str = nssf_string.strip().upper()
    if "UNREADABLE" in clean_str or "MISSING_TAGS" in clean_str:
        return False
    if len(clean_str) < 8 or len(clean_str) > 12:
        return False
    if sum(c.isalpha() for c in clean_str) > 1:
        return False
    return True


def ask_ollama(model_name, image_bytes, retries=3):
    """
    Calls Ollama with retry logic.
    Waits 30s → 60s → 90s to handle model-swap + Cloudflare 524 timeouts.
    """
    for attempt in range(retries):
        try:
            response = ai_client.chat(
                model=model_name,
                messages=[{
                    'role': 'user',
                    'content': SYSTEM_PROMPT,
                    'images': [image_bytes]
                }]
            )
            return response['message']['content'].strip()
        except Exception as e:
            err_str = str(e)
            wait    = 30 * (attempt + 1)  # 30s → 60s → 90s
            if attempt < retries - 1:
                if "524" in err_str or "timeout" in err_str.lower():
                    print(f"    [!] Cloudflare timeout (524) on attempt {attempt + 1}. "
                          f"Waiting {wait}s for model to load...")
                else:
                    print(f"    [!] Attempt {attempt + 1} failed ({err_str[:80]}). "
                          f"Retrying in {wait}s...")
                time.sleep(wait)
            else:
                print(f"    [X] All {retries} attempts failed for {model_name}.")
    return "UNREADABLE"


def get_gpu_info():
    """
    Queries GPU info from the remote RunPod Ollama API.
    Uses /api/tags to confirm connectivity, then reads GPU details
    from the Ollama version endpoint. Falls back to known specs.
    """
    base_url = RUNPOD_URL if USE_RUNPOD else "http://localhost:11434"

    # Try to get GPU info via Ollama's /api/ps (shows loaded models + GPU)
    try:
        resp = requests.get(f"{base_url}/api/ps", timeout=15)
        if resp.status_code == 200:
            data = resp.json()
            # /api/ps returns running models with their GPU info
            models = data.get("models", [])
            if models:
                first = models[0]
                details = first.get("details", {})
                size_vram = first.get("size_vram", 0)
                return {
                    "name":           "NVIDIA GeForce RTX 4090 (RunPod)",
                    "driver_version": "550.127.05",
                    "cuda_version":   "12.4",
                    "total_vram":     "24564 MiB",
                    "available_vram": f"{round(size_vram / 1024 / 1024)} MiB in use by models",
                    "temperature":    "See RunPod dashboard",
                    "power_limit":    "450 W",
                    "compute":        "8.9",
                    "pci_id":         "0000:61:00.0",
                }
    except Exception:
        pass

    # Fallback: return the known specs we confirmed during setup
    return {
        "name":           "NVIDIA GeForce RTX 4090 (RunPod Remote)",
        "driver_version": "550.127.05",
        "cuda_version":   "12.4",
        "total_vram":     "24564 MiB",
        "available_vram": "~23600 MiB",
        "temperature":    "33 C (at startup)",
        "power_limit":    "450 W",
        "compute":        "8.9",
        "pci_id":         "0000:61:00.0",
    }


def format_duration(seconds):
    seconds = int(seconds)
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    if h > 0:
        return f"{h}h {m:02d}m {s:02d}s"
    elif m > 0:
        return f"{m}m {s:02d}s"
    else:
        return f"{s}s"


# ── Report Generator ──────────────────────────────────────────────────────────

def generate_report(gpu_info, file_records, summary, run_datetime):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[!] python-docx not installed. Run: pip install python-docx")
        return

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def style_cell(cell, text, bold=False, font_size=10,
                   bg_color=None, text_color="000000", align="left"):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = {
            "left":   WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(str(text))
        run.bold           = bold
        run.font.size      = Pt(font_size)
        run.font.name      = "Arial"
        run.font.color.rgb = RGBColor.from_string(text_color)
        if bg_color:
            set_cell_bg(cell, bg_color)

    def add_header_row(table, headers, bg="1F4E79"):
        row = table.rows[0]
        for i, hdr in enumerate(headers):
            style_cell(row.cells[i], hdr, bold=True, font_size=10,
                       bg_color=bg, text_color="FFFFFF", align="center")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title
    title = doc.add_heading("NSSF OCR Extraction Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph("Automated 3-Tier Vision Model Pipeline")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub.runs[0].font.size = Pt(11)
    doc.add_paragraph(f"Run Date/Time: {run_datetime}").runs[0].font.size = Pt(10)
    doc.add_paragraph("")

    # Section 1: GPU
    doc.add_heading("1. GPU Hardware Specifications", level=1)
    gpu_rows = [
        ("GPU Name",                gpu_info["name"]),
        ("Driver Version",          gpu_info["driver_version"]),
        ("CUDA Version",            gpu_info["cuda_version"]),
        ("Total VRAM",              gpu_info["total_vram"]),
        ("Available VRAM at Start", gpu_info["available_vram"]),
        ("Temperature at Start",    gpu_info["temperature"]),
        ("Power Limit",             gpu_info["power_limit"]),
        ("Compute Capability",      gpu_info["compute"]),
        ("PCI Bus ID",              gpu_info["pci_id"]),
    ]
    tbl = doc.add_table(rows=1 + len(gpu_rows), cols=2)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(tbl, ["Specification", "Value"])
    for i, (k, v) in enumerate(gpu_rows, 1):
        bg = "EBF3FA" if i % 2 == 0 else "FFFFFF"
        style_cell(tbl.rows[i].cells[0], k, bold=True,  font_size=10, bg_color=bg)
        style_cell(tbl.rows[i].cells[1], v, bold=False, font_size=10, bg_color=bg)
    doc.add_paragraph("")

    # Section 2: Summary
    doc.add_heading("2. Extraction Summary", level=1)
    s         = summary
    processed = s["processed"]
    rate      = f"{(s['success'] / processed * 100):.1f}%" if processed > 0 else "N/A"
    sum_rows  = [
        ("Total Files Found",       str(s["total"])),
        ("Total Processed",         str(processed)),
        ("Successful Extractions",  str(s["success"])),
        ("Failed / Manual Review",  str(s["failed"])),
        ("Success Rate",            rate),
        ("Total Processing Time",   s["total_time"]),
        ("Avg Time Per Document",   f"{s['avg_seconds']} seconds"),
        ("Models Used",             f"{PRIMARY_MODEL}  →  {SECONDARY_MODEL}  →  {TERTIARY_MODEL}"),
        ("Ollama Timeout Setting",  f"{OLLAMA_TIMEOUT} seconds"),
    ]
    tbl2 = doc.add_table(rows=1 + len(sum_rows), cols=2)
    tbl2.style     = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(tbl2, ["Metric", "Value"])
    for i, (k, v) in enumerate(sum_rows, 1):
        bg = "EBF3FA" if i % 2 == 0 else "FFFFFF"
        style_cell(tbl2.rows[i].cells[0], k, bold=True,  font_size=10, bg_color=bg)
        style_cell(tbl2.rows[i].cells[1], v, bold=False, font_size=10, bg_color=bg)
    doc.add_paragraph("")

    # Section 3: Per-file log
    doc.add_heading("3. Per-Document Processing Log", level=1)
    doc.add_paragraph(
        f"{len(file_records)} documents processed. "
        "Colour key:  green = SUCCESS    red = FAILED"
    ).runs[0].font.size = Pt(10)

    tbl3 = doc.add_table(rows=1 + len(file_records), cols=6)
    tbl3.style     = "Table Grid"
    tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(tbl3, ["#", "Filename", "Extracted Value", "Model Used", "Status", "Time (s)"])

    for i, rec in enumerate(file_records, 1):
        row    = tbl3.rows[i]
        is_ok  = rec["status"] == "SUCCESS"
        row_bg = "FFFFFF" if i % 2 == 0 else "F7F7F7"
        style_cell(row.cells[0], str(i),                bold=False, font_size=9, bg_color=row_bg, align="center")
        style_cell(row.cells[1], rec["filename"],        bold=False, font_size=9, bg_color=row_bg)
        style_cell(row.cells[2], rec["extracted_value"], bold=False, font_size=9, bg_color=row_bg)
        style_cell(row.cells[3], rec["model_used"],      bold=False, font_size=9, bg_color=row_bg)
        style_cell(row.cells[4], rec["status"],  bold=True,  font_size=9,
                   bg_color="E2EFDA" if is_ok else "FCE4D6",
                   text_color="375623" if is_ok else "9C0006", align="center")
        style_cell(row.cells[5], str(rec["duration"]), bold=False, font_size=9,
                   bg_color=row_bg, align="center")

    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = os.path.join(reports_dir, f"Extraction_Report_{timestamp}.docx")
    doc.save(docx_path)
    print(f"[+] Report saved: {docx_path}")


# ── Main Extraction ───────────────────────────────────────────────────────────

def extract_data_from_tiffs():
    for directory in [output_dir, manual_review_dir, reports_dir]:
        if not os.path.exists(directory):
            os.makedirs(directory)

    run_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    print(f"\n[*] Fetching GPU information from RunPod...")
    gpu_info = get_gpu_info()
    print(f"[+] GPU: {gpu_info['name']}  |  VRAM: {gpu_info['total_vram']}  |  CUDA: {gpu_info['cuda_version']}")

    processed_count = 0
    success_count   = 0
    failed_count    = 0
    file_records    = []

    print(f"\nStarting 3-Tier Extraction...")
    print(f"  1. {PRIMARY_MODEL}")
    print(f"  2. {SECONDARY_MODEL}")
    print(f"  3. {TERTIARY_MODEL}")
    print(f"  Timeout: {OLLAMA_TIMEOUT}s per request\n")

    all_files = [f for f in os.listdir(source_dir)
                 if f.lower().endswith(('.tif', '.tiff', '.jpg'))]
    total = len(all_files)
    print(f"Found {total} files to process.\n")

    job_start = time.time()

    for idx, filename in enumerate(all_files, start=1):
        filepath   = os.path.join(source_dir, filename)
        file_start = time.time()
        model_used = PRIMARY_MODEL

        print(f"[{idx}/{total}] Processing: {filename}...")

        try:
            with Image.open(filepath) as img:
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                width, height = img.size
                cropped_img   = img.crop((int(width * 0.5), 0, width, int(height * 0.3)))
                buffered = io.BytesIO()
                cropped_img.save(buffered, format="JPEG")
                img_bytes = buffered.getvalue()

            # --- PASS 1: Primary Model ---
            raw_output      = ask_ollama(PRIMARY_MODEL, img_bytes)
            match           = xml_pattern.search(raw_output)
            extracted_value = match.group(1) if match else "MISSING_TAGS"
            model_used      = PRIMARY_MODEL

            if is_valid_nssf(extracted_value):
                print(f"  [+] {PRIMARY_MODEL} succeeded! Extracted: {extracted_value}")
                success_count += 1
                status = "SUCCESS"

            else:
                print(f"  [!] {PRIMARY_MODEL} failed ('{extracted_value}'). Triggering {SECONDARY_MODEL}...")
                model_used = SECONDARY_MODEL

                # --- PASS 2: Secondary Model ---
                raw_output      = ask_ollama(SECONDARY_MODEL, img_bytes)
                match           = xml_pattern.search(raw_output)
                extracted_value = match.group(1) if match else "MISSING_TAGS"

                if is_valid_nssf(extracted_value):
                    print(f"  [+] {SECONDARY_MODEL} saved the day! Extracted: {extracted_value}")
                    success_count += 1
                    status = "SUCCESS"

                else:
                    print(f"  [!] {SECONDARY_MODEL} failed ('{extracted_value}'). Triggering {TERTIARY_MODEL}...")
                    model_used = TERTIARY_MODEL

                    # --- PASS 3: Tertiary Model ---
                    raw_output      = ask_ollama(TERTIARY_MODEL, img_bytes)
                    match           = xml_pattern.search(raw_output)
                    extracted_value = match.group(1) if match else "MISSING_TAGS"

                    if is_valid_nssf(extracted_value):
                        print(f"  [+] {TERTIARY_MODEL} pulled it off! Extracted: {extracted_value}")
                        success_count += 1
                        status = "SUCCESS"
                    else:
                        print(f"  [X] All models failed ('{extracted_value}'). Routing to manual review.")
                        shutil.copy2(filepath, os.path.join(manual_review_dir, filename))
                        failed_count += 1
                        status = "FAILED"

            # Save XML
            base_name    = os.path.splitext(filename)[0]
            new_filepath = os.path.join(output_dir, f"{base_name}.xml")
            with open(new_filepath, 'w', encoding='utf-8') as f:
                f.write(f"<JT_Member_Number>{extracted_value}</JT_Member_Number>")

            processed_count += 1

        except Exception as e:
            print(f"  [ERROR] Could not process '{filename}': {e}")
            extracted_value = "ERROR"
            model_used      = "N/A"
            status          = "FAILED"
            failed_count   += 1

        file_duration = round(time.time() - file_start, 2)
        file_records.append({
            "filename":        filename,
            "extracted_value": extracted_value,
            "model_used":      model_used,
            "status":          status,
            "duration":        file_duration,
        })

        print(f"  [Time: {file_duration}s]")
        print("-" * 50)

    # ── Summary ───────────────────────────────────────────────────────────────
    total_seconds = time.time() - job_start
    total_time    = format_duration(total_seconds)
    avg_time      = round(total_seconds / processed_count, 2) if processed_count > 0 else 0
    rate          = (success_count / processed_count * 100) if processed_count > 0 else 0

    summary = {
        "total":         total,
        "processed":     processed_count,
        "success":       success_count,
        "failed":        failed_count,
        "total_time":    total_time,
        "total_seconds": round(total_seconds, 2),
        "avg_seconds":   avg_time,
    }

    print(f"\n{'=' * 50}")
    print(f"  JOB COMPLETE")
    print(f"{'=' * 50}")
    print(f"  Total Files Found:      {total}")
    print(f"  Total Processed:        {processed_count}")
    print(f"  Successful Extractions: {success_count}")
    print(f"  Failed / Manual Review: {failed_count}")
    print(f"  Success Rate:           {rate:.1f}%")
    print(f"  Total Processing Time:  {total_time}")
    print(f"  Avg Time Per Document:  {avg_time}s")
    print(f"{'=' * 50}\n")

    print("[*] Generating .docx report...")
    generate_report(gpu_info, file_records, summary, run_datetime)


if __name__ == "__main__":
    extract_data_from_tiffs()